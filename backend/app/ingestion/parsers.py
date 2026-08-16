"""
Document parsers — one function per supported file type, all normalized
to the same output shape: a list of ParsedPage, each carrying the text
content and (when meaningful) a page number for citation purposes.

OCR-ready architecture: PDFParser first tries native text extraction; if
a page yields near-zero text (common for scanned/image-only pages), it's
flagged via `needs_ocr` so an OCR backend (e.g. pytesseract) can be
plugged in later without changing the pipeline's public interface.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

import markdown as markdown_lib
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models.document import DocumentType


@dataclass
class ParsedPage:
    text: str
    page_number: int | None
    needs_ocr: bool = False


class ParsingError(Exception):
    """Raised when a document cannot be parsed."""


MIN_CHARS_BEFORE_OCR_FLAG = 20


@dataclass
class DocumentMetadata:
    title: str | None = None
    author: str | None = None
    created_at: str | None = None  # ISO string; source formats vary too much to normalize further here


def extract_pdf_metadata(file_bytes: bytes) -> DocumentMetadata:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        info = reader.metadata or {}
    except Exception:  # noqa: BLE001 — metadata is best-effort, never fatal
        return DocumentMetadata()

    return DocumentMetadata(
        title=getattr(info, "title", None),
        author=getattr(info, "author", None),
        created_at=str(getattr(info, "creation_date", None)) if getattr(info, "creation_date", None) else None,
    )


def extract_docx_metadata(file_bytes: bytes) -> DocumentMetadata:
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        props = doc.core_properties
    except Exception:  # noqa: BLE001
        return DocumentMetadata()

    return DocumentMetadata(
        title=props.title or None,
        author=props.author or None,
        created_at=props.created.isoformat() if props.created else None,
    )


_METADATA_EXTRACTORS = {
    DocumentType.PDF: extract_pdf_metadata,
    DocumentType.DOCX: extract_docx_metadata,
}


def extract_document_metadata(file_bytes: bytes, document_type: DocumentType) -> DocumentMetadata:
    extractor = _METADATA_EXTRACTORS.get(document_type)
    if extractor is None:
        return DocumentMetadata()  # TXT/Markdown carry no structured metadata
    return extractor(file_bytes)


def parse_pdf(file_bytes: bytes) -> list[ParsedPage]:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001
        raise ParsingError(f"Failed to open PDF: {exc}") from exc

    pages: list[ParsedPage] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        pages.append(
            ParsedPage(
                text=text,
                page_number=i,
                needs_ocr=len(text) < MIN_CHARS_BEFORE_OCR_FLAG,
            )
        )
    return pages


def parse_docx(file_bytes: bytes) -> list[ParsedPage]:
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001
        raise ParsingError(f"Failed to open DOCX: {exc}") from exc

    # DOCX has no native page boundaries — treat the whole document as one
    # logical unit; the chunker downstream splits it further.
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [ParsedPage(text=text, page_number=None)]


def parse_txt(file_bytes: bytes) -> list[ParsedPage]:
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")
    return [ParsedPage(text=text, page_number=None)]


def parse_markdown(file_bytes: bytes) -> list[ParsedPage]:
    raw = file_bytes.decode("utf-8")
    html = markdown_lib.markdown(raw)
    # Store the original markdown as chunk text (preserves structure/links
    # for the LLM) rather than the rendered HTML.
    return [ParsedPage(text=raw, page_number=None)]


_PARSERS = {
    DocumentType.PDF: parse_pdf,
    DocumentType.DOCX: parse_docx,
    DocumentType.TXT: parse_txt,
    DocumentType.MARKDOWN: parse_markdown,
}

_EXTENSION_TO_TYPE = {
    ".pdf": DocumentType.PDF,
    ".docx": DocumentType.DOCX,
    ".txt": DocumentType.TXT,
    ".md": DocumentType.MARKDOWN,
    ".markdown": DocumentType.MARKDOWN,
}


def document_type_from_filename(filename: str) -> DocumentType:
    ext = Path(filename).suffix.lower()
    if ext not in _EXTENSION_TO_TYPE:
        raise ParsingError(f"Unsupported file extension: {ext}")
    return _EXTENSION_TO_TYPE[ext]


def parse_document(file_bytes: bytes, document_type: DocumentType) -> list[ParsedPage]:
    parser = _PARSERS.get(document_type)
    if parser is None:
        raise ParsingError(f"No parser registered for {document_type}")
    return parser(file_bytes)
