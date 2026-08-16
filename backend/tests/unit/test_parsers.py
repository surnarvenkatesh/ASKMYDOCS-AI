"""
Unit tests for app.ingestion.parsers.
"""
import io

import pytest

from app.ingestion.parsers import (
    ParsingError,
    document_type_from_filename,
    extract_docx_metadata,
    extract_pdf_metadata,
    parse_docx,
    parse_markdown,
    parse_pdf,
    parse_txt,
)
from app.models.document import DocumentType


@pytest.mark.unit
class TestDocumentTypeFromFilename:
    @pytest.mark.parametrize(
        "filename,expected",
        [
            ("report.pdf", DocumentType.PDF),
            ("Report.PDF", DocumentType.PDF),
            ("notes.docx", DocumentType.DOCX),
            ("readme.txt", DocumentType.TXT),
            ("readme.md", DocumentType.MARKDOWN),
            ("readme.markdown", DocumentType.MARKDOWN),
        ],
    )
    def test_recognized_extensions(self, filename, expected):
        assert document_type_from_filename(filename) == expected

    def test_unsupported_extension_raises(self):
        with pytest.raises(ParsingError):
            document_type_from_filename("archive.zip")


@pytest.mark.unit
class TestParseTxt:
    def test_parses_utf8_text(self):
        pages = parse_txt("Hello, world!".encode("utf-8"))
        assert len(pages) == 1
        assert pages[0].text == "Hello, world!"
        assert pages[0].page_number is None

    def test_falls_back_to_latin1_on_bad_utf8(self):
        raw = "café".encode("latin-1")
        pages = parse_txt(raw)
        assert "caf" in pages[0].text


@pytest.mark.unit
class TestParseMarkdown:
    def test_preserves_raw_markdown(self):
        md = "# Title\n\nSome **bold** text."
        pages = parse_markdown(md.encode("utf-8"))
        assert len(pages) == 1
        assert pages[0].text == md


@pytest.mark.unit
class TestParseDocx:
    def test_parses_paragraphs(self):
        from docx import Document as DocxDocument

        buffer = io.BytesIO()
        doc = DocxDocument()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.save(buffer)

        pages = parse_docx(buffer.getvalue())
        assert len(pages) == 1
        assert "First paragraph." in pages[0].text
        assert "Second paragraph." in pages[0].text

    def test_invalid_docx_raises_parsing_error(self):
        with pytest.raises(ParsingError):
            parse_docx(b"not a real docx file")


@pytest.mark.unit
class TestParsePdf:
    def test_parses_single_page_with_text(self):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        buffer = io.BytesIO()
        writer.write(buffer)

        pages = parse_pdf(buffer.getvalue())
        assert len(pages) == 1
        assert pages[0].page_number == 1
        # Blank page -> no extractable text -> should be flagged for OCR
        assert pages[0].needs_ocr is True

    def test_invalid_pdf_raises_parsing_error(self):
        with pytest.raises(ParsingError):
            parse_pdf(b"not a real pdf file")


@pytest.mark.unit
class TestExtractDocxMetadata:
    def test_extracts_title_and_author(self):
        from docx import Document as DocxDocument

        buffer = io.BytesIO()
        doc = DocxDocument()
        doc.core_properties.title = "Q3 Board Deck"
        doc.core_properties.author = "Jane Doe"
        doc.add_paragraph("Body text.")
        doc.save(buffer)

        metadata = extract_docx_metadata(buffer.getvalue())
        assert metadata.title == "Q3 Board Deck"
        assert metadata.author == "Jane Doe"

    def test_corrupt_file_returns_empty_metadata_without_raising(self):
        metadata = extract_docx_metadata(b"not a docx")
        assert metadata.title is None
        assert metadata.author is None


@pytest.mark.unit
class TestExtractPdfMetadata:
    def test_blank_pdf_returns_metadata_object_without_raising(self):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        buffer = io.BytesIO()
        writer.write(buffer)

        metadata = extract_pdf_metadata(buffer.getvalue())
        assert metadata.title is None or isinstance(metadata.title, str)

    def test_corrupt_file_returns_empty_metadata_without_raising(self):
        metadata = extract_pdf_metadata(b"not a pdf")
        assert metadata.title is None
        assert metadata.author is None
